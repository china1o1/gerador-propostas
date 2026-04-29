import streamlit as st
from docx import Document
import io
from datetime import date, timedelta
import random

# Função para gerar as datas automaticamente
def gerar_datas():
    hoje = date.today()
    data_emissao = hoje.strftime("%d/%m/%Y")
    ultimo_dia_ano = date(hoje.year, 12, 31)
    dias_restantes = (ultimo_dia_ano - hoje).days
    dias_aleatorios = random.randint(15, dias_restantes) if dias_restantes >= 15 else 15
    data_fim = hoje + timedelta(days=dias_aleatorios)
    data_vigencia = f"{data_emissao} - {data_fim.strftime('%d/%m/%Y')}"
    return data_emissao, data_vigencia
# --- NOVAS FUNÇÕES PARA CALCULAR DINHEIRO ---
def converter_para_numero(valor_str):
    if not valor_str: return 0.0
    texto_limpo = valor_str.replace("R$", "").replace(".", "").replace(",", ".").strip()
    try:
        return float(texto_limpo)
    except:
        return 0.0

def formatar_moeda(valor_float):
    return f"{valor_float:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
# --------------------------------------------
# --- FUNÇÃO PARA DELETAR LINHAS DA TABELA ---
def remover_linhas_extras(doc):
    for tabela in doc.tables:
        for row in tabela.rows:
            deletar = False
            for cell in row.cells:
                if "DELETAR_LINHA" in cell.text:
                    deletar = True
                    break
            if deletar:
                tbl = row._tr.getparent()
                if row._tr in tbl:
                    tbl.remove(row._tr)
# --------------------------------------------

# Função que procura as TAGS no Word (nas linhas, tabelas e cabeçalhos)
def substituir_texto(doc, dicionario_dados):
    # 1. Substituir no corpo do texto (parágrafos e tabelas)
    for paragrafo in doc.paragraphs:
        for tag, valor in dicionario_dados.items():
            if tag in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(tag, str(valor))
                
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for tag, valor in dicionario_dados.items():
                        if tag in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(tag, str(valor))

    # 2. Varredura completa em TODOS os tipos de cabeçalho
    for secao in doc.sections:
        # Lista todos os cabeçalhos possíveis (Principal, Primeira Página, Páginas Pares)
        headers = [secao.header, secao.first_page_header, secao.even_page_header]
        
        for header in headers:
            # Procura em parágrafos do cabeçalho
            for paragrafo in header.paragraphs:
                for tag, valor in dicionario_dados.items():
                    if tag in paragrafo.text:
                        paragrafo.text = paragrafo.text.replace(tag, str(valor))
            
            # Procura em tabelas dentro do cabeçalho
            for tabela in header.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            for tag, valor in dicionario_dados.items():
                                if tag in paragrafo.text:
                                    paragrafo.text = paragrafo.text.replace(tag, str(valor))
# ==========================================
# Função ESPECÍFICA para buscar dentro de Caixas de Texto (Shapes)
def substituir_em_shapes(doc, dicionario_dados):
    for shape in doc.shapes:
        if shape.has_text_frame:
            for paragrafo in shape.text_frame.paragraphs:
                for tag, valor in dicionario_dados.items():
                    if tag in paragrafo.text:
                        paragrafo.text = paragrafo.text.replace(tag, str(valor))
# INTERFACE DO SITE
# ==========================================
st.set_page_config(page_title="Gerador de Propostas SENAI", page_icon="📄", layout="wide")

st.title("📄 Gerador de Propostas - SENAI")
st.write("Preencha os dados abaixo para gerar o documento Word formatado automaticamente.")

st.subheader("1. Dados do Cliente")
col1, col2 = st.columns(2)

with col1:
    nome_empresa = st.text_input("Nome da Empresa Solicitante")
    cnpj = st.text_input("CNPJ")
    categoria = st.selectbox("Categoria", ["INDUSTRIA", "COMUNIDADE"])
    endereco = st.text_input("Endereço de Execução")

with col2:
    contato_funcao = st.text_input("Contato/Função (Ex: João - Gerente)")
    telefone = st.text_input("Telefone de Contato")
    email = st.text_input("E-mail")
    num_pessoas = st.text_input("Nº de pessoas atendidas")
    responsavel = st.text_input("Responsável pela Proposta", value="ELITON GABRIEL SILVA CORDEIRO")

st.subheader("2. Dados do Serviço e Valores")
# Pergunta quantos serviços a pessoa quer (de 1 a 10)
num_servicos = st.number_input("Quantos serviços deseja adicionar à proposta?", min_value=1, max_value=10, value=1)

servicos_lista = []
# Cria os campos na tela repetidas vezes com base no número escolhido
for i in range(int(num_servicos)):
    st.markdown(f"**🔹 Serviço {i+1}**")
    col3, col4 = st.columns(2)
    with col3:
        s_nome = st.text_input(f"Serviço {i+1}", key=f"s_nome_{i}")
        s_desc = st.text_area(f"Descrição do Serviço {i+1}", key=f"s_desc_{i}")
        s_unid = st.text_input(f"Unidade Executora {i+1}", key=f"s_unid_{i}")
    with col4:
        s_qtd = st.text_input(f"Quantidade {i+1}", key=f"s_qtd_{i}")
        # Pede apenas os números. O código vai adicionar o R$ depois.
        s_val_un = st.text_input(f"Valor Unitário {i+1} (Ex: 400,00)", key=f"s_val_un_{i}")
        s_val_tot = st.text_input(f"Valor Total {i+1} (Ex: 6.000,00)", key=f"s_val_tot_{i}")

    # Salva os dados desse serviço em uma lista
    servicos_lista.append({
        "nome": s_nome, "desc": s_desc, "unid": s_unid,
        "qtd": s_qtd, "val_un": s_val_un, "val_tot": s_val_tot
    })

# --- CAMPO DO DESCONTO ---
st.subheader("3. Resumo e Desconto")
desconto_input = st.text_input("Valor do Desconto em Reais (Ex: 500,00) - Deixe em branco se não houver")
# -------------------------

st.write("---")

# Botão principal
if st.button("Gerar Proposta"):
    if not nome_empresa:
        st.warning("⚠️ Por favor, preencha pelo menos o nome da empresa.")
    else:
        try:
            # Carrega o documento Word que serve de molde
            doc = Document("Template_Proposta.docx")
            
            # Gera as datas
            data_emissao, data_vigencia = gerar_datas()
            
            # ==========================================
            # 1. A MÁQUINA DE CALCULAR (Soma e Subtração)
            # ==========================================
            soma_global = 0.0
            
            # O Python olha cada serviço preenchido e soma os valores totais
            for serv in servicos_lista:
                soma_global += converter_para_numero(serv["val_tot"])
            
            # Pega o desconto digitado. Se estiver vazio, a função transforma em 0.0
            valor_desconto = converter_para_numero(desconto_input)
            
            # A SUBTRAÇÃO ACONTECE AQUI:
            valor_final = soma_global - valor_desconto
            
            # ==========================================
            # 2. TRANSFORMANDO DE VOLTA EM TEXTO PARA O WORD
            # ==========================================
            texto_global = f"R$ {formatar_moeda(soma_global)}"
            texto_desconto = f"R$ {formatar_moeda(valor_desconto)}" if valor_desconto > 0 else "-"
            texto_final = f"R$ {formatar_moeda(valor_final)}"

            # ==========================================
            # 3. O DICIONÁRIO BASE ATUALIZADO
            # ==========================================
            dados_para_trocar = {
                "{{EMPRESA}}": nome_empresa,
                "{{CNPJ}}": cnpj,
                "{{CATEGORIA}}": categoria,
                "{{CONTATO_FUNCAO}}": contato_funcao,
                "{{ENDERECO}}": endereco,
                "{{TELEFONE}}": telefone,
                "{{EMAIL}}": email,
                "{{NUM_PESSOAS}}": num_pessoas,
                "{{DATA_EMISSAO}}": data_emissao,
                "{{DATA_VIGENCIA}}": data_vigencia,
                "{{RESPONSAVEL}}": responsavel,
                "{{VALOR_GLOBAL}}": texto_global,  # <-- Entra o valor somado
                "{{DESCONTO}}": texto_desconto,    # <-- Entra o desconto
                "{{VALOR_FINAL}}": texto_final     # <-- Entra o resultado da subtração
            }

            # 2. Preenche os serviços dinamicamente (Varre do 1 ao 10)
            for i in range(10):
                idx = i + 1
                if i < num_servicos:
                    # Se o serviço foi preenchido no site, coloca na tag correspondente
                    dados_para_trocar[f"{{{{SERVICO_{idx}}}}}"] = servicos_lista[i]["nome"]
                    dados_para_trocar[f"{{{{DESCRICAO_{idx}}}}}"] = servicos_lista[i]["desc"]
                    dados_para_trocar[f"{{{{UNIDADE_{idx}}}}}"] = servicos_lista[i]["unid"]
                    dados_para_trocar[f"{{{{QTD_{idx}}}}}"] = servicos_lista[i]["qtd"]
                    # Aqui o Python embute o "R$" direto no valor!
                    dados_para_trocar[f"{{{{VALOR_UN_{idx}}}}}"] = f"R$ {servicos_lista[i]['val_un']}" if servicos_lista[i]['val_un'] else ""
                    dados_para_trocar[f"{{{{VALOR_TOTAL_{idx}}}}}"] = f"R$ {servicos_lista[i]['val_tot']}" if servicos_lista[i]['val_tot'] else ""
                else:
                    # Se o serviço NÃO foi usado, injetamos a senha para deletar a linha
                    dados_para_trocar[f"{{{{SERVICO_{idx}}}}}"] = "DELETAR_LINHA"
                    dados_para_trocar[f"{{{{DESCRICAO_{idx}}}}}"] = ""
                    dados_para_trocar[f"{{{{UNIDADE_{idx}}}}}"] = ""
                    dados_para_trocar[f"{{{{QTD_{idx}}}}}"] = ""
                    dados_para_trocar[f"{{{{VALOR_UN_{idx}}}}}"] = ""
                    dados_para_trocar[f"{{{{VALOR_TOTAL_{idx}}}}}"] = ""
            
            # Executa a varredura e substituição no documento
            substituir_texto(doc, dados_para_trocar)

            # --- NOVO: ACIONA A DESTRUIÇÃO DAS LINHAS EXTRAS ---
            remover_linhas_extras(doc)
            # ---------------------------------------------------
            
            # Salva o documento pronto na "memória" do site para download
            arquivo_memoria = io.BytesIO()
            doc.save(arquivo_memoria)
            arquivo_memoria.seek(0)
            
            st.success("✨ Proposta processada com sucesso!")
            
            # Cria o botão de download para o usuário baixar
            st.download_button(
                label="📥 Baixar Documento Word Pronto",
                data=arquivo_memoria,
                file_name=f"Proposta_{nome_empresa}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error("Erro ao gerar. Verifique se o arquivo 'Template_Proposta.docx' está na mesma pasta e fechado.")