import random
from datetime import date, timedelta
from docx import Document

def gerar_datas():
    """Gera a data de emissão (hoje) e uma vigência aleatória até o fim do ano."""
    hoje = date.today()
    data_emissao = hoje.strftime("%d/%m/%Y")
    
    # Calcula quantos dias faltam até o final do ano atual
    ultimo_dia_ano = date(hoje.year, 12, 31)
    dias_restantes = (ultimo_dia_ano - hoje).days
    
    # Gera uma data aleatória entre hoje e o fim do ano
    dias_aleatorios = random.randint(15, dias_restantes) # Pelo menos 15 dias de vigência
    data_fim = hoje + timedelta(days=dias_aleatorios)
    data_vigencia = f"{data_emissao} - {data_fim.strftime('%d/%m/%Y')}"
    
    return data_emissao, data_vigencia

def substituir_texto(doc, dicionario_dados):
    """Procura as tags nos parágrafos e nas tabelas e substitui pelos dados."""
    # 1. Substituir nos parágrafos normais
    for paragrafo in doc.paragraphs:
        for tag, valor in dicionario_dados.items():
            if tag in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(tag, str(valor))
                
    # 2. Substituir nas tabelas (Onde a maior parte da proposta do SENAI fica)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for tag, valor in dicionario_dados.items():
                        if tag in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(tag, str(valor))

def automatizar_proposta():
    # Caminho do seu template (deve estar na mesma pasta do script)
    caminho_template = "Template_Proposta.docx"
    doc = Document(caminho_template)
    
    # Gerando as datas automaticamente
    data_emissao, data_vigencia = gerar_datas()
    
    # Dicionário com os dados da NOVA proposta (Você só precisa alterar isso a cada proposta!)
    dados_nova_proposta = {
        "{{RESPONSAVEL}}": "ELITON GABRIEL SILVA CORDEIRO",
        "{{EMPRESA}}": "Indústria Exemplo S.A.",
        "{{CNPJ}}": "12.345.678/0001-99",
        "{{EMAIL}}": "contato@industriaexemplo.com.br",
        "{{TELEFONE}}": "81 9 9999-0000",
        "{{NUM_PESSOAS}}": "15",
        "{{ENDERECO}}": "Av. Principal, 1000 - Distrito Industrial",
        "{{SERVICO}}": "Curso de Excel Avançado e Power BI",
        "{{DESCRICAO}}": "Capacitar colaboradores na análise de dados e criação de dashboards gerenciais.",
        "{{UNIDADE}}": "SENAI - Santo Amaro",
        "{{VALOR_UN}}": "R$ 400,00",
        "{{QTD}}": "15",
        "{{VALOR_TOTAL}}": "R$ 6.000,00",
        "{{DATA_EMISSAO}}": data_emissao,
        "{{DATA_VIGENCIA}}": data_vigencia
        "{{NUMERO_DA_PROPOSTA}}": "PRO-210/2026"
    }
    
    # Roda a função de substituição
    substituir_texto(doc, dados_nova_proposta)
    
    # Salva o novo arquivo com o nome da empresa
    nome_arquivo_saida = f"Proposta_{dados_nova_proposta['{{EMPRESA}}']}.docx"
    doc.save(nome_arquivo_saida)
    print(f"Sucesso! A proposta '{nome_arquivo_saida}' foi gerada.")

# Executa o script
if __name__ == "__main__":
    automatizar_proposta()